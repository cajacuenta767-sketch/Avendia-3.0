import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document('c:/Users/PC/Documents/ChatGPT/Avend Escala 3.0/frontend/scratch/test_pca_fixed.docx')

found = False
for i, t in enumerate(doc.tables):
    for r in t.rows:
        for c in r.cells:
            if 'Acciones del Docente' in c.text:
                found = True
                print(f"=== Found 6.2 Table (Table Index {i}) ===")
                for r_idx, row in enumerate(t.rows):
                    print(f"Row {r_idx}:")
                    for c_idx, cell in enumerate(row.cells):
                        print(f"  Cell {c_idx} has {len(cell.paragraphs)} paragraphs:")
                        for p_idx, p in enumerate(cell.paragraphs):
                            print(f"    [P{p_idx}] {p.text}")
                break
        if found:
            break
    if found:
        break
