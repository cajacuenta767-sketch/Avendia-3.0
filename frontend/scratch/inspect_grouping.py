import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document('c:/Users/PC/Documents/ChatGPT/Avend Escala 3.0/frontend/scratch/test_grouping_fixed.docx')
print("Total tables in test_grouping_fixed.docx:", len(doc.tables))
for i, t in enumerate(doc.tables):
    print(f"\n--- Table {i} ({len(t.rows)} rows x {len(t.columns)} cols) ---")
    for r_idx, r in enumerate(t.rows):
        cells = [c.text.strip().replace('\n', ' | ') for c in r.cells]
        print(f"Row {r_idx}: {cells}")
