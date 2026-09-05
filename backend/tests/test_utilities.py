from uuid import uuid4

from httpx import ASGITransport, AsyncClient
from sqlalchemy import select
from test_community import teacher_token

from app.db.session import session_factory
from app.main import app
from app.modules.users.model import User


async def test_referral_credit_is_transactional_and_never_repeated():
    from app.modules.utilities.referrals import Referral, attribute_referral

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        admin = await headers(client, "rewards-admin@example.edu", True)
        owner = await headers(client, "referrer@example.edu")
        await headers(client, "invitee@example.edu")
        assert (await client.put("/api/v1/referrals/code", headers=owner)).status_code == 409
        assert (
            await client.put(
                "/api/v1/admin/referrals/settings",
                headers=owner,
                json={"enabled": True, "reward": 100},
            )
        ).status_code == 403
        assert (
            await client.put(
                "/api/v1/admin/referrals/settings",
                headers=admin,
                json={"enabled": True, "reward": 100},
            )
        ).status_code == 200
        code = (await client.put("/api/v1/referrals/code", headers=owner)).json()["code"]
        assert (await client.put("/api/v1/referrals/code", headers=owner)).json()["code"] == code
        async with session_factory() as db:
            invitee = await db.scalar(select(User).where(User.email == "invitee@example.edu"))
            await attribute_referral(db, invitee, code)
            await db.commit()
            referral = await db.scalar(select(Referral).where(Referral.invitee_id == invitee.id))
            rid = str(referral.id)
        before = (await client.get("/api/v1/referrals/me", headers=owner)).json()["balance"]
        decision = {"status": "credited", "reason": "Registro verificado en prueba"}
        assert (
            await client.post(f"/api/v1/admin/referrals/{rid}/review", headers=owner, json=decision)
        ).status_code == 403
        assert (
            await client.post(f"/api/v1/admin/referrals/{rid}/review", headers=admin, json=decision)
        ).status_code == 200
        assert (
            await client.post(f"/api/v1/admin/referrals/{rid}/review", headers=admin, json=decision)
        ).status_code == 409
        after = (await client.get("/api/v1/referrals/me", headers=owner)).json()
        assert after["balance"] == before + 100 and after["credited"] == 100
        assert "invitee_id" not in after["items"][0]


async def test_personal_utility_summary_is_private_and_uses_real_counts():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await headers(client, "summary-owner@example.edu")
        other = await headers(client, "summary-other@example.edu")
        await client.post(
            "/api/v1/documents",
            headers=owner,
            json={"title": "Plan del docente", "document_type": "plan", "content": "Contenido"},
        )
        await client.post(
            "/api/v1/community/posts",
            headers=owner,
            json={
                "title": "Aporte de aula",
                "content": "Comparto una estrategia que funcionó con mi grupo.",
                "education_level": "Primaria",
                "curricular_area": "Comunicación",
            },
        )
        summary = await client.get("/api/v1/utilities/summary", headers=owner)
        assert summary.status_code == 200, summary.text
        assert summary.json()["history"]["documents"] == 1
        assert summary.json()["community"]["posts"] == 1
        other_summary = (await client.get("/api/v1/utilities/summary", headers=other)).json()
        assert other_summary["history"]["documents"] == 0
        assert other_summary["community"]["posts"] == 0


async def test_document_revisions_conflicts_trash_and_access():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await headers(client, "historian@example.edu")
        other = await headers(client, "stranger@example.edu")
        doc = (
            await client.post(
                "/api/v1/documents",
                headers=owner,
                json={
                    "title": "Plan original",
                    "document_type": "plan",
                    "content": "Contenido original",
                },
            )
        ).json()
        path = f"/api/v1/documents/{doc['id']}"
        for field in ("title", "document_type", "status", "favorite", "metadata"):
            assert (await client.patch(path, headers=owner, json={field: None})).status_code == 422
        assert (await client.get(path, headers=owner)).json()["revision"] == 1
        changed = await client.patch(
            path,
            headers=owner,
            json={"title": "Plan corregido", "favorite": True, "expected_revision": 1},
        )
        assert changed.status_code == 200, changed.text
        assert changed.json()["revision"] == 2
        assert (
            await client.patch(
                path, headers=owner, json={"title": "Edición obsoleta", "expected_revision": 1}
            )
        ).status_code == 409
        versions = (await client.get(path + "/versions", headers=owner)).json()
        assert versions[0]["title"] == "Plan original"
        assert (await client.get(path + "/versions", headers=other)).status_code == 404
        restored = await client.post(path + f"/versions/{versions[0]['id']}/restore", headers=owner)
        assert restored.json()["revision"] == 3 and restored.json()["title"] == "Plan original"
        assert (await client.get("/api/v1/history?favorite=true", headers=owner)).json()[
            "total"
        ] == 1
        await client.delete(path, headers=owner)
        assert (await client.get(path, headers=owner)).status_code == 404
        assert (await client.get("/api/v1/history?trashed=true", headers=owner)).json()[
            "total"
        ] == 1
        assert (await client.get("/api/v1/history?trashed=true", headers=other)).json()[
            "total"
        ] == 0
        assert (await client.post(path + "/recover", headers=other)).status_code == 404
        assert (await client.post(path + "/recover", headers=owner)).status_code == 200
        assert (await client.get(path, headers=owner)).json()["title"] == "Plan original"
        feed = await client.get("/api/v1/history/feed?q=Plan&page=1&size=1", headers=owner)
        assert feed.status_code == 200, feed.text
        assert feed.json()["total"] == 1 and len(feed.json()["documents"]) == 1
        assert (await client.get("/api/v1/history/feed", headers=other)).json()["total"] == 0


async def test_community_unique_reactions_private_saved_and_comments():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await headers(client, "postowner@example.edu")
        other = await headers(client, "postreader@example.edu")
        payload = {
            "title": "Lecturas en comunidad",
            "content": "Organizamos una lectura con familias y docentes.",
            "education_level": "Primaria",
            "curricular_area": "Comunicación",
        }
        pid = (await client.post("/api/v1/community/posts", headers=owner, json=payload)).json()[
            "id"
        ]
        for _ in range(2):
            assert (
                await client.put(f"/api/v1/community/posts/{pid}/reactions/useful", headers=other)
            ).status_code == 200
        feed = (await client.get("/api/v1/community/feed", headers=other)).json()
        assert feed["items"][0]["useful_count"] == 1
        await client.put(f"/api/v1/community/posts/{pid}/reactions/saved", headers=other)
        assert (await client.get("/api/v1/community/feed?saved=true", headers=owner)).json()[
            "total"
        ] == 0
        assert (await client.get("/api/v1/community/feed?saved=true", headers=other)).json()[
            "total"
        ] == 1
        comment = {"content": "Gracias por compartir esta experiencia", "request_id": str(uuid4())}
        for _ in range(2):
            assert (
                await client.post(
                    f"/api/v1/community/posts/{pid}/comments", headers=other, json=comment
                )
            ).status_code == 201
        assert (await client.get(f"/api/v1/community/posts/{pid}/comments", headers=owner)).json()[
            "total"
        ] == 1
        await client.delete(f"/api/v1/community/posts/{pid}", headers=owner)
        assert (await client.get("/api/v1/community/feed", headers=other)).json()["total"] == 0
        assert (
            await client.get(f"/api/v1/community/posts/{pid}/comments", headers=other)
        ).status_code == 404


async def test_template_details_analysis_are_private_and_persistent():
    from io import BytesIO

    from docx import Document

    output = BytesIO()
    document = Document()
    document.add_paragraph("{{titulo}}")
    document.save(output)
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await headers(client, "templatesowner@example.edu")
        other = await headers(client, "templatesother@example.edu")
        uploaded = await client.post(
            "/api/v1/templates",
            headers=owner,
            files={
                "file": (
                    "base.docx",
                    output.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        tid = uploaded.json()["id"]
        path = f"/api/v1/templates/{tid}"
        data = {
            "name": "Plan institucional",
            "category": "pca",
            "description": "Formato de mi institución",
            "tags": ["2026"],
        }
        assert (await client.put(path + "/details", headers=owner, json=data)).status_code == 200
        assert (await client.get(path + "/details", headers=other)).status_code == 404
        result = await client.post(path + "/analyze", headers=owner)
        assert result.status_code == 200, result.text
        assert result.json()["fields"] == ["titulo"]
        persisted = (await client.get(path + "/details", headers=owner)).json()
        assert persisted["category"] == "pca" and persisted["analysis"]["sha256"]
        assert (await client.post(path + "/analyze", headers=other)).status_code == 404


async def test_template_versions_restore_original_bytes_and_recover():
    from io import BytesIO

    from docx import Document

    def make_file(text):
        output = BytesIO()
        document = Document()
        document.add_paragraph(text)
        document.save(output)
        return output.getvalue()

    original, replacement = make_file("Original institucional"), make_file("Nuevo formato")
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await headers(client, "versionowner@example.edu")
        other = await headers(client, "versionother@example.edu")
        uploaded = await client.post(
            "/api/v1/templates", headers=owner, files={"file": ("original.docx", original)}
        )
        path = f"/api/v1/templates/{uploaded.json()['id']}"
        changed = await client.post(
            path + "/replace",
            headers=owner,
            data={"expected_revision": "1"},
            files={"file": ("new.docx", replacement)},
        )
        assert changed.status_code == 200, changed.text
        assert changed.json()["revision"] == 2
        conflict = await client.post(
            path + "/replace",
            headers=owner,
            data={"expected_revision": "1"},
            files={"file": ("stale.docx", original)},
        )
        assert conflict.status_code == 409
        version = (await client.get(path + "/versions", headers=owner)).json()[0]
        assert (
            await client.get(path + f"/versions/{version['id']}/download", headers=owner)
        ).content == original
        assert (
            await client.get(path + f"/versions/{version['id']}/download", headers=other)
        ).status_code == 404
        restored = await client.post(path + f"/versions/{version['id']}/restore", headers=owner)
        assert restored.status_code == 200 and restored.json()["revision"] == 3
        assert (await client.get(path + "/download", headers=owner)).content == original
        await client.delete(path, headers=owner)
        assert (await client.get(path + "/download", headers=owner)).status_code == 404
        assert len((await client.get("/api/v1/templates?trashed=true", headers=owner)).json()) == 1
        assert (await client.post(path + "/recover", headers=other)).status_code == 404
        assert (await client.post(path + "/recover", headers=owner)).status_code == 200
        assert (await client.get(path + "/download", headers=owner)).content == original


async def test_admin_overview_contains_aggregates_not_private_content():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        teacher = await headers(client, "metricsteacher@example.edu")
        admin = await headers(client, "metricsadmin@example.edu", True)
        assert (
            await client.get("/api/v1/admin/utilities/overview", headers=teacher)
        ).status_code == 403
        await client.post(
            "/api/v1/documents",
            headers=teacher,
            json={
                "title": "Documento privado",
                "document_type": "plan",
                "content": "Contenido reservado",
            },
        )
        result = await client.get("/api/v1/admin/utilities/overview", headers=admin)
        assert result.status_code == 200 and result.json()["documents"] == 1
        assert "Contenido reservado" not in result.text


async def test_publication_retries_do_not_duplicate_records():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        admin = await headers(client, "idempotent@example.edu", True)
        tutorial = {
            "request_id": str(uuid4()),
            "title": "Tutorial verificable",
            "url": "https://example.com/video.mp4",
            "category": "Planificación",
        }
        first = await client.post("/api/v1/admin/tutorials", headers=admin, json=tutorial)
        second = await client.post("/api/v1/admin/tutorials", headers=admin, json=tutorial)
        assert first.status_code == 201 and second.status_code == 201
        assert first.json()["id"] == second.json()["id"]
        post = {
            "request_id": str(uuid4()),
            "title": "Publicación verificable",
            "content": "Un recurso para docentes de primaria",
            "education_level": "Primaria",
            "curricular_area": "Comunicación",
        }
        first = await client.post("/api/v1/community/posts", headers=admin, json=post)
        second = await client.post("/api/v1/community/posts", headers=admin, json=post)
        assert first.status_code == 201 and second.status_code == 201
        assert first.json()["id"] == second.json()["id"]


async def headers(client, email, admin=False):
    token = await teacher_token(client, email)
    if admin:
        async with session_factory() as db:
            user = await db.scalar(select(User).where(User.email == email))
            user.role = "admin"
            await db.commit()
    return {"Authorization": f"Bearer {token}"}


async def test_ideas_persist_votes_are_unique_permissions_and_notifications():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await headers(client, "owner@example.edu")
        other = await headers(client, "other@example.edu")
        admin = await headers(client, "admin@example.edu", True)
        payload = {
            "request_id": str(uuid4()),
            "title": "Mejorar exportaciones",
            "description": "La tabla debe conservar el ancho de columnas.",
            "category": "exportacion",
        }
        created = await client.post("/api/v1/ideas", headers=owner, json=payload)
        assert created.status_code == 201, created.text
        idea_id = created.json()["id"]
        assert (await client.post("/api/v1/ideas", headers=owner, json=payload)).json()[
            "id"
        ] == idea_id
        assert (await client.get("/api/v1/ideas", headers=owner)).json()["total"] == 1
        for _ in range(2):
            vote = await client.put(f"/api/v1/ideas/{idea_id}/vote", headers=other)
            assert vote.json()["votes"] == 1
        assert (
            await client.put(f"/api/v1/ideas/{idea_id}/vote?enabled=false", headers=other)
        ).json()["votes"] == 0
        assert (
            await client.patch(
                f"/api/v1/ideas/{idea_id}",
                headers=other,
                json={"title": "Otro título", "description": payload["description"]},
            )
        ).status_code == 403
        review = {"status": "planned", "response": "Validaremos la exportación de tablas."}
        assert (
            await client.patch(f"/api/v1/admin/ideas/{idea_id}", headers=owner, json=review)
        ).status_code == 403
        assert (
            await client.patch(f"/api/v1/admin/ideas/{idea_id}", headers=admin, json=review)
        ).status_code == 200
        assert (await client.get("/api/v1/notifications", headers=owner)).json()["unread"] == 1
        assert (await client.get("/api/v1/notifications", headers=other)).json()["unread"] == 0
        assert (await client.put("/api/v1/notifications/read", headers=owner)).status_code == 200
        assert (await client.get("/api/v1/notifications", headers=owner)).json()["unread"] == 0
        comment = {"request_id": str(uuid4()), "content": "También ocurre en mi formato."}
        for _ in range(2):
            assert (
                await client.post(f"/api/v1/ideas/{idea_id}/comments", headers=other, json=comment)
            ).status_code == 201
        assert (await client.get(f"/api/v1/ideas/{idea_id}/comments", headers=owner)).json()[
            "total"
        ] == 1
        assert (await client.get("/api/v1/ideas?page=0", headers=owner)).status_code == 422
        review["status"] = "hidden"
        await client.patch(f"/api/v1/admin/ideas/{idea_id}", headers=admin, json=review)
        assert (await client.get("/api/v1/ideas", headers=other)).json()["total"] == 0
        assert (
            await client.get(f"/api/v1/ideas/{idea_id}/comments", headers=other)
        ).status_code == 404


async def test_tutorial_publishing_and_individual_progress():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        admin = await headers(client, "editor@example.edu", True)
        teacher = await headers(client, "viewer@example.edu")
        other = await headers(client, "viewer2@example.edu")
        video = {
            "title": "Cómo crear una rúbrica",
            "url": "https://example.com/video.mp4",
            "category": "Evaluación",
            "transcript": "Explicación del criterio de evaluación",
        }
        assert (
            await client.post("/api/v1/admin/tutorials", headers=teacher, json=video)
        ).status_code == 403
        created = await client.post("/api/v1/admin/tutorials", headers=admin, json=video)
        assert created.status_code == 201, created.text
        tid = created.json()["id"]
        assert (await client.get("/api/v1/tutorials", headers=teacher)).json()["total"] == 0
        assert (
            await client.put(
                f"/api/v1/tutorials/{tid}/progress", headers=teacher, json={"completed": True}
            )
        ).status_code == 404
        video["published"] = True
        assert (
            await client.put(f"/api/v1/admin/tutorials/{tid}", headers=admin, json=video)
        ).status_code == 200
        assert (
            await client.put(
                f"/api/v1/tutorials/{tid}/progress",
                headers=teacher,
                json={"seconds": 42, "favorite": True},
            )
        ).status_code == 200
        first = (await client.get("/api/v1/tutorials?q=criterio", headers=teacher)).json()["items"][
            0
        ]
        assert first["seconds"] == 42 and first["favorite"] is True
        second = (await client.get("/api/v1/tutorials", headers=other)).json()["items"][0]
        assert second["seconds"] == 0 and second["favorite"] is False
        assert (
            await client.put(
                f"/api/v1/tutorials/{tid}/progress", headers=teacher, json={"seconds": -1}
            )
        ).status_code == 422
        video["url"] = "javascript:alert(1)"
        assert (
            await client.post("/api/v1/admin/tutorials", headers=admin, json=video)
        ).status_code == 422
