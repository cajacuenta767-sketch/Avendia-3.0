from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient
from openpyxl import load_workbook
from reportlab.pdfgen import canvas

from app.main import app


def registration(email: str) -> dict[str, object]:
    return {
        "email": email,
        "full_name": "María Gómez",
        "password": "a-secure-password",
        "dre": "DRE Lima Metropolitana",
        "ugel": "UGEL 03",
        "school_name": "I.E. José María Arguedas",
        "director_name": "Elena Torres",
        "education_modality": "EBR",
        "education_level": "Secundaria",
        "grade": "3° de Secundaria",
        "section": "A",
        "curricular_area": "Comunicación",
        "school_year": 2026,
    }


async def authenticated_headers(client: AsyncClient, email: str) -> dict[str, str]:
    payload = registration(email)
    assert (await client.post("/api/v1/auth/register", json=payload)).status_code == 201
    login = await client.post(
        "/api/v1/auth/login",
        json={"email": email, "password": payload["password"]},
    )
    assert login.status_code == 200
    return {"Authorization": f"Bearer {login.json()['access_token']}"}


async def roster_and_students(
    client: AsyncClient,
    headers: dict[str, str],
    *,
    section: str = "A",
) -> tuple[dict[str, object], list[dict[str, object]]]:
    roster_response = await client.post(
        "/api/v1/rosters",
        headers=headers,
        json={
            "school_year": 2026,
            "institution_name": "I.E. José María Arguedas",
            "modality": "EBR",
            "education_level": "Secundaria",
            "grade": "3° de Secundaria",
            "section": section,
            "name": f"Tercero {section}",
        },
    )
    assert roster_response.status_code == 201, roster_response.text
    roster = roster_response.json()
    students: list[dict[str, object]] = []
    for index, name in enumerate(("Ana Torres Salazar", "Luis Quispe Rojas"), 1):
        response = await client.post(
            f"/api/v1/rosters/{roster['id']}/students",
            headers=headers,
            json={"full_name": name, "internal_code": f"EST-{section}-{index:03d}"},
        )
        assert response.status_code == 201, response.text
        students.append(response.json())
    return roster, students


def checklist_payload(
    roster: dict[str, object],
    students: list[dict[str, object]],
    *,
    title: str = "Lista de cotejo de lectura **comprensiva**",
) -> dict[str, object]:
    return {
        "kind": "checklist",
        "title": title,
        "roster_id": roster["id"],
        "status": "draft",
        "general_data": {
            "teacher_name": "María Gómez",
            "institution_name": "I.E. José María Arguedas",
            "modality": "EBR",
            "education_level": "Secundaria",
            "grade": "3° de Secundaria",
            "curricular_area": "Comunicación",
            "activity": "Lectura argumentativa",
            "date": "2026-09-01",
            "period": "III bimestre",
        },
        "settings": {"scale": "yes_no_in_progress"},
        "general_observation": "Registro del aula.",
        "participants": [
            {
                "student_id": student["id"],
                "sort_order": index,
                "individual_notes": "Observación **sin Markdown**" if index == 0 else None,
            }
            for index, student in enumerate(students)
        ],
        "criteria": [
            {
                "client_key": "criterion-1",
                "code": "C1",
                "title": "Ubica información explícita",
                "description": "Identifica datos relevantes del texto.",
                "sort_order": 0,
            },
            {
                "client_key": "criterion-2",
                "code": "C2",
                "title": "Infiere información",
                "description": "Relaciona indicios para construir una conclusión.",
                "sort_order": 1,
            },
        ],
        "records": [
            {
                "student_id": students[0]["id"],
                "criterion_key": "criterion-1",
                "value": "yes",
                "observation": "Logro sostenido.",
            },
            {
                "student_id": students[0]["id"],
                "criterion_key": "criterion-2",
                "value": "in_progress",
            },
            {
                "student_id": students[1]["id"],
                "criterion_key": "criterion-1",
                "value": "no",
            },
            {
                "student_id": students[1]["id"],
                "criterion_key": "criterion-2",
                "value": "yes",
            },
        ],
        "observations": [],
    }


def pdf_file(text: str) -> bytes:
    output = BytesIO()
    document = canvas.Canvas(output)
    document.drawString(72, 760, text)
    document.save()
    return output.getvalue()


def docx_file(text: str) -> bytes:
    output = BytesIO()
    document = Document()
    document.add_heading("Lectura para analizar", level=1)
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Concepto"
    table.cell(0, 1).text = "Explicación"
    document.save(output)
    return output.getvalue()


@pytest.mark.asyncio
async def test_checklist_composite_draft_authorization_and_atomic_updates() -> None:
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await authenticated_headers(client, "evaluation-owner@example.edu")
        other = await authenticated_headers(client, "evaluation-other@example.edu")
        roster, students = await roster_and_students(client, owner)
        other_roster, other_students = await roster_and_students(client, other, section="B")

        created = await client.post(
            "/api/v1/evaluation-instruments",
            headers=owner,
            json=checklist_payload(roster, students),
        )
        assert created.status_code == 201, created.text
        body = created.json()
        instrument_id = body["id"]
        assert body["revision"] == 1
        assert body["status"] == "draft"
        assert [item["student_name"] for item in body["participants"]] == [
            "Ana Torres Salazar",
            "Luis Quispe Rojas",
        ]
        assert [item["criterion_key"] for item in body["records"]].count("criterion-1") == 2

        listing = await client.get(
            "/api/v1/evaluation-instruments?kind=checklist",
            headers=owner,
        )
        assert listing.status_code == 200
        assert listing.json()["items"][0]["participant_count"] == 2
        assert listing.json()["items"][0]["criterion_count"] == 2
        student_history = await client.get(
            f"/api/v1/evaluation-instruments?student_id={students[0]['id']}",
            headers=owner,
        )
        assert student_history.status_code == 200
        assert student_history.json()["total"] == 1

        for method, path in (
            (client.get, f"/api/v1/evaluation-instruments/{instrument_id}"),
            (client.get, f"/api/v1/evaluation-instruments/{instrument_id}/draft"),
            (client.get, f"/api/v1/evaluation-instruments/{instrument_id}/sources"),
            (client.get, f"/api/v1/evaluation-instruments/{instrument_id}/exports/checklist.xlsx"),
        ):
            assert (await method(path, headers=other)).status_code == 404

        stale_payload = checklist_payload(roster, students, title="Cambio que no debe guardarse")
        stale_payload["expected_revision"] = 0
        conflict = await client.put(
            f"/api/v1/evaluation-instruments/{instrument_id}/draft",
            headers=owner,
            json=stale_payload,
        )
        assert conflict.status_code == 409
        assert conflict.json()["detail"]["current_revision"] == 1
        unchanged = await client.get(
            f"/api/v1/evaluation-instruments/{instrument_id}", headers=owner
        )
        assert unchanged.json()["title"].startswith("Lista de cotejo")

        foreign_payload = checklist_payload(roster, students)
        foreign_payload["expected_revision"] = 1
        foreign_payload["participants"] = [
            {"student_id": students[0]["id"]},
            {"student_id": other_students[0]["id"]},
        ]
        foreign_payload["records"] = []
        rejected = await client.put(
            f"/api/v1/evaluation-instruments/{instrument_id}/draft",
            headers=owner,
            json=foreign_payload,
        )
        assert rejected.status_code == 422
        still_intact = await client.get(
            f"/api/v1/evaluation-instruments/{instrument_id}", headers=owner
        )
        assert still_intact.json()["revision"] == 1
        assert len(still_intact.json()["records"]) == 4

        valid_payload = checklist_payload(roster, students, title="Lista actualizada")
        valid_payload["expected_revision"] = 1
        valid_payload["status"] = "generated"
        updated = await client.put(
            f"/api/v1/evaluation-instruments/{instrument_id}/draft",
            headers=owner,
            json=valid_payload,
        )
        assert updated.status_code == 200, updated.text
        assert updated.json()["revision"] == 2
        assert updated.json()["status"] == "generated"

        assert (
            await client.delete(f"/api/v1/evaluation-instruments/{instrument_id}", headers=owner)
        ).status_code == 204
        assert (await client.get("/api/v1/evaluation-instruments", headers=owner)).json()[
            "total"
        ] == 0
        archived = await client.get(
            "/api/v1/evaluation-instruments?include_archived=true", headers=owner
        )
        assert archived.json()["items"][0]["status"] == "archived"
        restored = await client.post(
            f"/api/v1/evaluation-instruments/{instrument_id}/restore", headers=owner
        )
        assert restored.status_code == 200
        assert restored.json()["status"] == "draft"
        assert other_roster["id"] != roster["id"]


@pytest.mark.asyncio
async def test_rubric_feedback_and_observation_contracts_are_persisted() -> None:
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        headers = await authenticated_headers(client, "evaluation-rubric@example.edu")
        roster, students = await roster_and_students(client, headers)
        payload = {
            "kind": "rubric",
            "title": "Rúbrica analítica de exposición",
            "roster_id": roster["id"],
            "general_data": {"rubric_type": "analytic"},
            "settings": {"scale": "AD_A_B_C"},
            "participants": [{"student_id": item["id"]} for item in students],
            "criteria": [
                {
                    "client_key": "oral-expression",
                    "code": "C1",
                    "title": "Expresión oral",
                    "weight": 100,
                    "levels": [
                        {
                            "client_key": key,
                            "code": code,
                            "label": label,
                            "score": score,
                            "sort_order": index,
                        }
                        for index, (key, code, label, score) in enumerate(
                            (
                                ("ad", "AD", "Logro destacado", 4),
                                ("a", "A", "Logro esperado", 3),
                                ("b", "B", "En proceso", 2),
                                ("c", "C", "En inicio", 1),
                            )
                        )
                    ],
                }
            ],
            "records": [
                {
                    "student_id": students[0]["id"],
                    "criterion_key": "oral-expression",
                    "level_key": "a",
                    "evidence": "Presentación grabada.",
                    "strength": "Explica sus ideas con orden.",
                    "improvement": "Necesita sostener contacto visual.",
                    "recommendation": "Ensaya frente a un compañero y verifica tres pausas.",
                    "teacher_decision": "Se mantiene la calificación A.",
                }
            ],
            "observations": [
                {
                    "student_id": students[0]["id"],
                    "observed_at": "2026-09-01T14:30:00Z",
                    "situation": "Exposición de aula",
                    "focus": "Comunicación oral",
                    "objective_facts": "Expuso durante cuatro minutos y respondió dos preguntas.",
                    "context_factors": "Trabajo en equipo previo.",
                    "interpretation": "Organiza el contenido con autonomía.",
                    "conclusion": "Alcanza el logro esperado.",
                    "commitments": "Practicar contacto visual.",
                },
                {
                    "student_id": None,
                    "observed_at": "2026-09-01T14:30:00Z",
                    "situation": "Exposición de aula",
                    "focus": "Participación grupal",
                    "objective_facts": "El equipo respetó los turnos.",
                    "common_to_group": True,
                },
            ],
        }
        created = await client.post("/api/v1/evaluation-instruments", headers=headers, json=payload)
        assert created.status_code == 201, created.text
        body = created.json()
        assert body["criteria"][0]["levels"][1]["code"] == "A"
        assert body["records"][0]["level_key"] == "a"
        assert body["records"][0]["recommendation"].startswith("Ensaya")
        assert body["records"][0]["teacher_decision"].startswith("Se mantiene")
        assert body["observations"][1]["student_id"] is None
        assert body["observations"][1]["common_to_group"] is True
        wrong_export = await client.get(
            f"/api/v1/evaluation-instruments/{body['id']}/exports/checklist.xlsx",
            headers=headers,
        )
        assert wrong_export.status_code == 409


@pytest.mark.asyncio
async def test_generated_evaluation_instruments_reject_semantically_incomplete_products() -> None:
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        headers = await authenticated_headers(client, "evaluation-quality@example.edu")
        roster, students = await roster_and_students(client, headers)

        incomplete_checklist = checklist_payload(roster, students)
        incomplete_checklist["status"] = "generated"
        incomplete_checklist["records"][0]["value"] = None
        checklist_response = await client.post(
            "/api/v1/evaluation-instruments",
            headers=headers,
            json=incomplete_checklist,
        )
        assert checklist_response.status_code == 422
        assert "una respuesta por estudiante e indicador" in checklist_response.text

        incomplete_rubric = {
            "kind": "rubric",
            "status": "generated",
            "title": "Rúbrica incompleta",
            "roster_id": roster["id"],
            "participants": [{"student_id": student["id"]} for student in students],
            "criteria": [
                {
                    "client_key": "criterion-1",
                    "code": "C1",
                    "title": "Criterio único",
                    "levels": [],
                }
            ],
            "records": [],
        }
        rubric_response = await client.post(
            "/api/v1/evaluation-instruments",
            headers=headers,
            json=incomplete_rubric,
        )
        assert rubric_response.status_code == 422
        assert "entre 3 y 6 criterios" in rubric_response.text

        incomplete_observation = {
            "kind": "observation",
            "status": "generated",
            "title": "Ficha de observación incompleta",
            "roster_id": roster["id"],
            "settings": {"mode": "individual", "scale_type": "Descriptiva"},
            "participants": [{"student_id": students[0]["id"]}],
            "criteria": [
                {
                    "client_key": "criterion-1",
                    "code": "C1",
                    "title": "Explica su estrategia con evidencias",
                }
            ],
            "observations": [
                {
                    "student_id": students[0]["id"],
                    "observed_at": "2026-09-04T14:30:00Z",
                    "situation": "Trabajo cooperativo",
                    "focus": "Explicación de estrategias",
                    "objective_facts": "Explicó el procedimiento y comparó dos resultados.",
                }
            ],
        }
        observation_response = await client.post(
            "/api/v1/evaluation-instruments",
            headers=headers,
            json=incomplete_observation,
        )
        assert observation_response.status_code == 422
        assert "interpretación, conclusión y compromisos" in observation_response.text

        source_document_without_artifact = {
            "kind": "text_questions",
            "status": "generated",
            "title": "Preguntas sobre el ciclo del agua",
            "general_data": {
                "source": {"pasted_text": "El agua se evapora por efecto del calor.", "sources": []}
            },
            "settings": {"literal_count": 1, "inferential_count": 1, "critical_count": 1},
        }
        missing_artifact_response = await client.post(
            "/api/v1/evaluation-instruments",
            headers=headers,
            json=source_document_without_artifact,
        )
        assert missing_artifact_response.status_code == 422
        assert "Genera y revisa" in missing_artifact_response.text

        source_document_without_artifact["settings"]["generated_artifact"] = {
            "document_title": "Preguntas sobre el ciclo del agua",
            "quality_status": "ready",
            "sections": [
                {"title": "Preguntas literales", "key_points": ["¿Qué causa la evaporación?"]},
                {"title": "Preguntas inferenciales", "key_points": ["¿Qué ocurriría sin calor?"]},
                {
                    "title": "Preguntas crítico-reflexivas",
                    "key_points": ["¿Cómo cuidarías el agua?"],
                },
                {
                    "title": "Respuestas esperadas",
                    "key_points": ["El calor", "No se evaporaría", "Respuesta argumentada"],
                },
            ],
        }
        complete_source_response = await client.post(
            "/api/v1/evaluation-instruments",
            headers=headers,
            json=source_document_without_artifact,
        )
        assert complete_source_response.status_code == 201, complete_source_response.text
        assert complete_source_response.json()["status"] == "generated"


@pytest.mark.asyncio
async def test_checklist_xlsx_contains_matrix_and_criterion_definitions() -> None:
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        headers = await authenticated_headers(client, "evaluation-excel@example.edu")
        roster, students = await roster_and_students(client, headers)
        created = await client.post(
            "/api/v1/evaluation-instruments",
            headers=headers,
            json=checklist_payload(roster, students),
        )
        instrument_id = created.json()["id"]
        response = await client.get(
            f"/api/v1/evaluation-instruments/{instrument_id}/exports/checklist.xlsx",
            headers=headers,
        )
        assert response.status_code == 200, response.text
        assert response.content.startswith(b"PK")
        workbook = load_workbook(BytesIO(response.content))
        assert workbook.sheetnames == ["Lista de cotejo", "Criterios"]
        matrix = workbook["Lista de cotejo"]
        assert matrix["A1"].value == "Lista de cotejo de lectura comprensiva"
        assert [matrix.cell(4, column).value for column in range(1, 7)] == [
            "N.°",
            "Estudiante",
            "Código",
            "C1",
            "C2",
            "Observación",
        ]
        assert [matrix.cell(5, column).value for column in range(2, 6)] == [
            "Ana Torres Salazar",
            "EST-A-001",
            "Sí",
            "En proceso",
        ]
        assert [matrix.cell(6, column).value for column in range(2, 6)] == [
            "Luis Quispe Rojas",
            "EST-A-002",
            "No",
            "Sí",
        ]
        assert "*" not in matrix["F5"].value
        criteria = workbook["Criterios"]
        assert criteria["A2"].value == "C1"
        assert criteria["C2"].value == "Identifica datos relevantes del texto."
        workbook.close()


@pytest.mark.asyncio
async def test_checklist_xlsx_neutralizes_formulas_and_reads_legacy_general_keys() -> None:
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        headers = await authenticated_headers(client, "evaluation-excel-security@example.edu")
        roster, students = await roster_and_students(client, headers)
        malicious_student = await client.patch(
            f"/api/v1/rosters/{roster['id']}/students/{students[0]['id']}",
            headers=headers,
            json={"full_name": "=1+1", "internal_code": "+EST-001"},
        )
        assert malicious_student.status_code == 200, malicious_student.text

        payload = checklist_payload(roster, students, title="=1+1")
        payload["general_data"] = {
            "teacherName": "\t=DOCENTE()",
            "directorName": "@DIRECTOR",
            "institution": "@INSTITUCIÓN",
            "modality": "EBR",
            "level": "-Secundaria",
            "grade": "3° de Secundaria",
            "area": "+Comunicación",
            "activity": "=ACTIVIDAD()",
            "date": "2026-09-01",
            "period": "III bimestre",
        }
        payload["participants"][0]["individual_notes"] = "\r\n@NOTA"
        payload["criteria"][0].update(
            {"code": "@C1", "title": "=2+2", "description": "\t+SUM(1,1)"}
        )
        payload["records"][0]["observation"] = "\n-CMD"

        created = await client.post("/api/v1/evaluation-instruments", headers=headers, json=payload)
        assert created.status_code == 201, created.text
        response = await client.get(
            f"/api/v1/evaluation-instruments/{created.json()['id']}/exports/checklist.xlsx",
            headers=headers,
        )
        assert response.status_code == 200, response.text

        workbook = load_workbook(BytesIO(response.content), data_only=False)
        matrix = workbook["Lista de cotejo"]
        definitions = workbook["Criterios"]
        assert matrix["A1"].value == "'=1+1"
        assert matrix["D4"].value == "'@C1"
        assert matrix["B5"].value == "'=1+1"
        assert matrix["C5"].value == "'+EST-001"
        assert "Docente: '=DOCENTE()" in matrix["A2"].value
        assert "Director(a): '@DIRECTOR" in matrix["A2"].value
        assert "Institución: '@INSTITUCIÓN" in matrix["A2"].value
        assert "Nivel: '-Secundaria" in matrix["A2"].value
        assert "Área: '+Comunicación" in matrix["A2"].value
        assert "Actividad / evidencia: '=ACTIVIDAD()" in matrix["A2"].value
        assert definitions["A2"].value == "'@C1"
        assert definitions["B2"].value == "'=2+2"
        assert definitions["C2"].value == "'+SUM(1,1)"
        assert "'-CMD" in matrix["F5"].value
        assert "'@NOTA" in matrix["F5"].value
        for worksheet in workbook.worksheets:
            assert all(cell.data_type != "f" for row in worksheet.iter_rows() for cell in row)
        workbook.close()


@pytest.mark.asyncio
async def test_real_pdf_docx_sources_security_and_owner_isolation() -> None:
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        owner = await authenticated_headers(client, "evaluation-files-owner@example.edu")
        other = await authenticated_headers(client, "evaluation-files-other@example.edu")
        created = await client.post(
            "/api/v1/evaluation-instruments",
            headers=owner,
            json={
                "kind": "text_questions",
                "title": "Preguntas sobre texto",
                "general_data": {"source_mode": "file"},
                "settings": {"reading_size": "large", "question_size": "medium"},
            },
        )
        assert created.status_code == 201, created.text
        instrument_id = created.json()["id"]

        pdf = pdf_file("La biodiversidad sostiene el equilibrio de los ecosistemas.")
        preview = await client.post(
            "/api/v1/evaluation-instruments/sources/extract",
            headers=owner,
            files={"file": ("lectura.pdf", pdf, "application/pdf")},
        )
        assert preview.status_code == 200, preview.text
        assert "biodiversidad" in preview.json()["extracted_text"]
        stored_pdf = await client.post(
            f"/api/v1/evaluation-instruments/{instrument_id}/sources",
            headers=owner,
            files={"file": ("lectura.pdf", pdf, "application/pdf")},
        )
        assert stored_pdf.status_code == 201, stored_pdf.text
        pdf_id = stored_pdf.json()["id"]

        duplicate = await client.post(
            f"/api/v1/evaluation-instruments/{instrument_id}/sources",
            headers=owner,
            files={"file": ("copia.pdf", pdf, "application/pdf")},
        )
        assert duplicate.status_code == 409
        assert (
            await client.get(
                f"/api/v1/evaluation-instruments/{instrument_id}/sources", headers=other
            )
        ).status_code == 404
        assert (
            await client.get(
                f"/api/v1/evaluation-instruments/{instrument_id}/sources/{pdf_id}/download",
                headers=other,
            )
        ).status_code == 404

        downloaded = await client.get(
            f"/api/v1/evaluation-instruments/{instrument_id}/sources/{pdf_id}/download",
            headers=owner,
        )
        assert downloaded.status_code == 200
        assert downloaded.content == pdf

        docx = docx_file("La lectura permite formular preguntas literales e inferenciales.")
        stored_docx = await client.post(
            f"/api/v1/evaluation-instruments/{instrument_id}/sources",
            headers=owner,
            files={
                "file": (
                    "lectura.docx",
                    docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert stored_docx.status_code == 201, stored_docx.text
        assert "inferenciales" in stored_docx.json()["extracted_text"]
        assert "Concepto | Explicación" in stored_docx.json()["extracted_text"]

        old_doc = await client.post(
            "/api/v1/evaluation-instruments/sources/extract",
            headers=owner,
            files={"file": ("antiguo.doc", b"old-word", "application/msword")},
        )
        assert old_doc.status_code == 415
        assert "DOC antiguo" in old_doc.text

        mismatch = await client.post(
            "/api/v1/evaluation-instruments/sources/extract",
            headers=owner,
            files={"file": ("falso.pdf", docx, "application/pdf")},
        )
        assert mismatch.status_code == 422

        macro_docx = BytesIO(docx)
        with ZipFile(macro_docx, mode="a", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/vbaProject.bin", b"not-a-real-macro")
        macro = await client.post(
            "/api/v1/evaluation-instruments/sources/extract",
            headers=owner,
            files={
                "file": (
                    "macro.docx",
                    macro_docx.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert macro.status_code == 422
        assert "macros" in macro.text.lower()

        oversized = await client.post(
            "/api/v1/evaluation-instruments/sources/extract",
            headers=owner,
            files={"file": ("grande.pdf", b"x" * (10 * 1024 * 1024 + 1), "application/pdf")},
        )
        assert oversized.status_code == 413

        assert (
            await client.delete(
                f"/api/v1/evaluation-instruments/{instrument_id}/sources/{pdf_id}",
                headers=owner,
            )
        ).status_code == 204
        remaining = await client.get(
            f"/api/v1/evaluation-instruments/{instrument_id}/sources", headers=owner
        )
        assert [item["extension"] for item in remaining.json()] == ["docx"]
