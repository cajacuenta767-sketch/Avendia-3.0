from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile, is_zipfile

from docx import Document
from fastapi import HTTPException, status
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.modules.evaluation_instruments.model import (
    EvaluationCriterion,
    EvaluationDraft,
    EvaluationInstrument,
    EvaluationLevel,
    EvaluationObservation,
    EvaluationParticipant,
    EvaluationRecord,
    EvaluationSourceFile,
)
from app.modules.evaluation_instruments.schemas import (
    CriterionRead,
    DraftWrite,
    InstrumentRead,
    InstrumentWrite,
    LevelRead,
    ObservationRead,
    ParticipantRead,
    RecordRead,
    SourceExtractionPreview,
    SourceRead,
)
from app.modules.rosters.model import Roster, Student
from app.modules.users.model import User

MAX_SOURCE_SIZE = 10 * 1024 * 1024
MAX_EXTRACTED_TEXT = 500_000
MAX_PDF_PAGES = 200
MAX_DOCX_UNCOMPRESSED_SIZE = 50 * 1024 * 1024

PDF_MIME_TYPES = {"application/pdf", "application/octet-stream", ""}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
    "application/octet-stream",
    "",
}


class SourceDocumentError(Exception):
    def __init__(self, message: str, status_code: int = 422) -> None:
        super().__init__(message)
        self.message = message
        self.status_code = status_code


@dataclass(frozen=True)
class ParsedSource:
    filename: str
    media_type: str
    extension: str
    byte_size: int
    sha256: str
    extracted_text: str
    content: bytes

    def preview(self) -> SourceExtractionPreview:
        return SourceExtractionPreview(
            filename=self.filename,
            media_type=self.media_type,
            extension=self.extension,
            byte_size=self.byte_size,
            sha256=self.sha256,
            extracted_text=self.extracted_text,
        )


def _clean_extracted_text(value: str) -> str:
    value = value.replace("\x00", "")
    value = re.sub(r"[\t ]+\n", "\n", value)
    value = re.sub(r"\n{4,}", "\n\n\n", value)
    value = value.strip()
    if len(value) > MAX_EXTRACTED_TEXT:
        raise SourceDocumentError(
            "El documento contiene demasiado texto para procesarlo de forma segura."
        )
    if not value:
        raise SourceDocumentError(
            "No se encontró texto legible. Si es un PDF escaneado, "
            "conviértelo con OCR antes de subirlo."
        )
    return value


def _extract_pdf(content: bytes) -> str:
    if not content.startswith(b"%PDF-"):
        raise SourceDocumentError("El contenido del archivo no corresponde a un PDF válido.")
    forbidden_markers = (
        b"/JavaScript",
        b"/JS ",
        b"/Launch",
        b"/EmbeddedFile",
        b"/OpenAction",
    )
    if any(marker in content for marker in forbidden_markers):
        raise SourceDocumentError(
            "El PDF contiene acciones, scripts o archivos incrustados y no puede procesarse."
        )
    try:
        reader = PdfReader(BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise SourceDocumentError("No se admiten archivos PDF protegidos con contraseña.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise SourceDocumentError(f"El PDF no puede superar {MAX_PDF_PAGES} páginas.")
        root = reader.trailer.get("/Root")
        if root is not None:
            root_object = root.get_object() if hasattr(root, "get_object") else root
            names = root_object.get("/Names") if hasattr(root_object, "get") else None
            if names is not None:
                names_object = names.get_object() if hasattr(names, "get_object") else names
                if names_object.get("/EmbeddedFiles") or names_object.get("/JavaScript"):
                    raise SourceDocumentError(
                        "El PDF contiene scripts o archivos incrustados y no puede procesarse."
                    )
            if root_object.get("/OpenAction") or root_object.get("/AA"):
                raise SourceDocumentError(
                    "El PDF contiene acciones automáticas y no puede procesarse."
                )
        pieces: list[str] = []
        size = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            size += len(text)
            if size > MAX_EXTRACTED_TEXT:
                raise SourceDocumentError(
                    "El documento contiene demasiado texto para procesarlo de forma segura."
                )
            pieces.append(text)
    except SourceDocumentError:
        raise
    except Exception as exc:
        raise SourceDocumentError("No se pudo leer el PDF; verifica que no esté dañado.") from exc
    return _clean_extracted_text("\n\n".join(pieces))


def _validate_docx_archive(content: bytes) -> None:
    if not is_zipfile(BytesIO(content)):
        raise SourceDocumentError("El contenido del archivo no corresponde a un DOCX válido.")
    try:
        with ZipFile(BytesIO(content)) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise SourceDocumentError("El archivo no contiene un documento Word DOCX válido.")
            unsafe_parts = (
                "vbaproject.bin",
                "word/embeddings/",
                "word/activex/",
                "word/oleobject",
            )
            lowered = [name.casefold() for name in names]
            if any(any(part in name for part in unsafe_parts) for name in lowered):
                raise SourceDocumentError(
                    "El Word contiene macros, objetos o archivos incrustados y no puede procesarse."
                )
            total_uncompressed = 0
            for info in archive.infolist():
                normalized = info.filename.replace("\\", "/")
                if normalized.startswith("/") or ".." in normalized.split("/"):
                    raise SourceDocumentError("El DOCX contiene rutas internas no seguras.")
                total_uncompressed += info.file_size
                if total_uncompressed > MAX_DOCX_UNCOMPRESSED_SIZE:
                    raise SourceDocumentError(
                        "El contenido descomprimido del DOCX excede el límite."
                    )
                if info.compress_size and info.file_size / info.compress_size > 250:
                    raise SourceDocumentError(
                        "El DOCX presenta una compresión anómala y fue rechazado."
                    )
    except SourceDocumentError:
        raise
    except BadZipFile as exc:
        raise SourceDocumentError("No se pudo abrir el DOCX; verifica que no esté dañado.") from exc


def _extract_docx(content: bytes) -> str:
    _validate_docx_archive(content)
    try:
        document = Document(BytesIO(content))
        pieces = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    pieces.append(row_text)
    except Exception as exc:
        raise SourceDocumentError("No se pudo extraer el texto del DOCX.") from exc
    return _clean_extracted_text("\n".join(pieces))


def parse_source_document(
    filename: str | None,
    content_type: str | None,
    content: bytes,
) -> ParsedSource:
    safe_name = Path(filename or "").name.strip()
    if not safe_name:
        raise SourceDocumentError("El archivo debe tener un nombre válido.")
    if not content:
        raise SourceDocumentError("El archivo está vacío.")
    if len(content) > MAX_SOURCE_SIZE:
        raise SourceDocumentError(
            "El archivo supera el límite de 10 MB.", status.HTTP_413_CONTENT_TOO_LARGE
        )

    extension = Path(safe_name).suffix.casefold()
    media_type = (content_type or "").split(";", 1)[0].strip().casefold()
    if extension == ".doc":
        raise SourceDocumentError(
            "El formato DOC antiguo no se admite de forma segura. "
            "Ábrelo en Word y guárdalo como DOCX.",
            status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        )
    if extension == ".pdf":
        if media_type not in PDF_MIME_TYPES:
            raise SourceDocumentError(
                "El tipo declarado no coincide con un PDF.",
                status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            )
        extracted_text = _extract_pdf(content)
        normalized_media_type = "application/pdf"
    elif extension == ".docx":
        if media_type not in DOCX_MIME_TYPES:
            raise SourceDocumentError(
                "El tipo declarado no coincide con un Word DOCX.",
                status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            )
        extracted_text = _extract_docx(content)
        normalized_media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        raise SourceDocumentError(
            "Formato no permitido. Sube un archivo PDF o Word DOCX.",
            status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        )
    return ParsedSource(
        filename=safe_name[:240],
        media_type=normalized_media_type,
        extension=extension.removeprefix("."),
        byte_size=len(content),
        sha256=hashlib.sha256(content).hexdigest(),
        extracted_text=extracted_text,
        content=content,
    )


def instrument_load_options() -> tuple[object, ...]:
    return (
        selectinload(EvaluationInstrument.participants).selectinload(EvaluationParticipant.student),
        selectinload(EvaluationInstrument.criteria).selectinload(EvaluationCriterion.levels),
        selectinload(EvaluationInstrument.records),
        selectinload(EvaluationInstrument.observations),
        selectinload(EvaluationInstrument.sources),
        selectinload(EvaluationInstrument.draft),
    )


async def owned_instrument(
    instrument_id: object,
    user: User,
    db: AsyncSession,
    *,
    for_update: bool = False,
    detailed: bool = False,
) -> EvaluationInstrument:
    query = select(EvaluationInstrument).where(
        EvaluationInstrument.id == instrument_id,
        EvaluationInstrument.owner_id == user.id,
    )
    if detailed:
        query = query.options(*instrument_load_options())
    if for_update:
        query = query.with_for_update()
    instrument = await db.scalar(query)
    if instrument is None:
        raise HTTPException(status_code=404, detail="Instrumento de evaluación no encontrado")
    return instrument


async def _validate_roster_and_students(
    payload: InstrumentWrite,
    user: User,
    db: AsyncSession,
) -> dict[object, Student]:
    if payload.roster_id is None:
        return {}
    roster = await db.scalar(
        select(Roster).where(Roster.id == payload.roster_id, Roster.owner_id == user.id)
    )
    if roster is None:
        raise HTTPException(status_code=404, detail="Nómina no encontrada")
    student_ids = {participant.student_id for participant in payload.participants}
    if not student_ids:
        return {}
    students = list(
        await db.scalars(
            select(Student).where(
                Student.roster_id == roster.id,
                Student.id.in_(student_ids),
            )
        )
    )
    if len(students) != len(student_ids):
        raise HTTPException(
            status_code=422,
            detail="Uno o más estudiantes no pertenecen a la nómina seleccionada.",
        )
    return {student.id: student for student in students}


async def replace_instrument_composite(
    instrument: EvaluationInstrument,
    payload: InstrumentWrite | DraftWrite,
    user: User,
    db: AsyncSession,
) -> None:
    if isinstance(payload, DraftWrite) and payload.expected_revision is not None:
        if payload.expected_revision != instrument.revision:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail={
                    "message": "El borrador fue modificado en otra sesión.",
                    "current_revision": instrument.revision,
                },
            )

    await _validate_roster_and_students(payload, user, db)

    await db.execute(
        delete(EvaluationObservation).where(EvaluationObservation.instrument_id == instrument.id)
    )
    await db.execute(
        delete(EvaluationRecord).where(EvaluationRecord.instrument_id == instrument.id)
    )
    await db.execute(delete(EvaluationLevel).where(EvaluationLevel.instrument_id == instrument.id))
    await db.execute(
        delete(EvaluationCriterion).where(EvaluationCriterion.instrument_id == instrument.id)
    )
    await db.execute(
        delete(EvaluationParticipant).where(EvaluationParticipant.instrument_id == instrument.id)
    )
    await db.flush()

    participants: dict[object, EvaluationParticipant] = {}
    for index, item in enumerate(payload.participants):
        sort_order = item.sort_order if "sort_order" in item.model_fields_set else index
        participant = EvaluationParticipant(
            instrument_id=instrument.id,
            student_id=item.student_id,
            role=item.role,
            team_name=item.team_name,
            sort_order=sort_order,
            common_notes=item.common_notes,
            individual_notes=item.individual_notes,
        )
        participants[item.student_id] = participant
        db.add(participant)

    criteria: dict[str, EvaluationCriterion] = {}
    level_inputs: list[tuple[EvaluationCriterion, object, int]] = []
    for index, item in enumerate(payload.criteria):
        sort_order = item.sort_order if "sort_order" in item.model_fields_set else index
        criterion = EvaluationCriterion(
            instrument_id=instrument.id,
            client_key=item.client_key,
            code=item.code,
            title=item.title,
            description=item.description,
            weight=item.weight,
            sort_order=sort_order,
        )
        criteria[item.client_key.casefold()] = criterion
        db.add(criterion)
        for level_index, level in enumerate(item.levels):
            level_inputs.append((criterion, level, level_index))
    await db.flush()

    levels: dict[tuple[str, str], EvaluationLevel] = {}
    for criterion, item, index in level_inputs:
        sort_order = item.sort_order if "sort_order" in item.model_fields_set else index
        level = EvaluationLevel(
            instrument_id=instrument.id,
            criterion_id=criterion.id,
            client_key=item.client_key,
            code=item.code,
            label=item.label,
            description=item.description,
            score=item.score,
            sort_order=sort_order,
        )
        levels[(criterion.client_key.casefold(), item.client_key.casefold())] = level
        db.add(level)
    await db.flush()

    for item in payload.records:
        criterion_key = item.criterion_key.casefold()
        criterion = criteria[criterion_key]
        level = (
            levels[(criterion_key, item.level_key.casefold())]
            if item.level_key is not None
            else None
        )
        db.add(
            EvaluationRecord(
                instrument_id=instrument.id,
                participant_id=participants[item.student_id].id,
                criterion_id=criterion.id,
                level_id=level.id if level else None,
                value=item.value,
                evidence=item.evidence,
                strength=item.strength,
                improvement=item.improvement,
                recommendation=item.recommendation,
                teacher_decision=item.teacher_decision,
                observation=item.observation,
            )
        )

    for item in payload.observations:
        participant = participants.get(item.student_id) if item.student_id else None
        db.add(
            EvaluationObservation(
                instrument_id=instrument.id,
                participant_id=participant.id if participant else None,
                observed_at=item.observed_at,
                situation=item.situation,
                focus=item.focus,
                objective_facts=item.objective_facts,
                context_factors=item.context_factors,
                interpretation=item.interpretation,
                conclusion=item.conclusion,
                commitments=item.commitments,
                common_to_group=item.common_to_group,
            )
        )

    stored_sources = list(
        await db.scalars(
            select(EvaluationSourceFile)
            .where(EvaluationSourceFile.instrument_id == instrument.id)
            .order_by(EvaluationSourceFile.created_at)
        )
    )
    canonical_general_data = _canonical_source_general_data(
        payload.general_data,
        stored_sources,
    )

    instrument.kind = payload.kind
    instrument.title = payload.title
    instrument.roster_id = payload.roster_id
    instrument.status = payload.status
    instrument.general_data_json = canonical_general_data
    instrument.settings_json = payload.settings
    instrument.general_observation = payload.general_observation
    instrument.archived_at = None
    instrument.revision += 1

    snapshot_payload = payload.model_dump(mode="json", exclude={"expected_revision"})
    snapshot_payload["general_data"] = canonical_general_data
    draft = await db.scalar(
        select(EvaluationDraft).where(EvaluationDraft.instrument_id == instrument.id)
    )
    if draft is None:
        draft = EvaluationDraft(
            instrument_id=instrument.id,
            owner_id=user.id,
            revision=instrument.revision,
            payload_json=snapshot_payload,
        )
        db.add(draft)
    else:
        draft.revision = instrument.revision
        draft.payload_json = snapshot_payload
    await db.flush()


def serialize_source(
    source: EvaluationSourceFile,
    *,
    instrument_revision: int,
) -> SourceRead:
    return SourceRead(
        id=source.id,
        filename=source.filename,
        media_type=source.media_type,
        extension=source.extension,
        byte_size=source.byte_size,
        sha256=source.sha256,
        extracted_text=source.extracted_text,
        extraction_status=source.extraction_status,
        created_at=source.created_at,
        instrument_revision=instrument_revision,
    )


def _canonical_source_general_data(
    general_data: dict[str, object],
    sources: list[EvaluationSourceFile],
) -> dict[str, object]:
    """Keep file references aligned with the owned source rows.

    Extracted text and binary content stay in their dedicated table.  The JSON draft only
    carries user-authored overrides, which avoids duplicating large documents and makes a
    deleted row impossible to retain as a ghost reference.
    """

    canonical = dict(general_data)
    raw_source = canonical.get("source")
    if not isinstance(raw_source, dict) and not sources:
        return canonical

    source_data = dict(raw_source) if isinstance(raw_source, dict) else {}
    raw_references = source_data.get("sources")
    references = raw_references if isinstance(raw_references, list) else []

    # Migrate the first frontend contract without trusting its filename or extracted text.
    legacy_id = source_data.get("source_id")
    if legacy_id:
        references = [
            *references,
            {
                "source_id": legacy_id,
                "edited_text": source_data.get("edited_text"),
            },
        ]

    reference_by_id = {
        str(item.get("source_id")): item
        for item in references
        if isinstance(item, dict) and item.get("source_id")
    }
    canonical_references: list[dict[str, object]] = []
    for source in sources:
        reference = reference_by_id.get(str(source.id), {})
        item: dict[str, object] = {
            "source_id": str(source.id),
            "filename": source.filename,
        }
        edited_text = reference.get("edited_text")
        if isinstance(edited_text, str) and edited_text != source.extracted_text:
            item["edited_text"] = edited_text
        canonical_references.append(item)

    for obsolete in ("source_id", "filename", "extracted_text", "edited_text"):
        source_data.pop(obsolete, None)
    source_data["sources"] = canonical_references
    canonical["source"] = source_data
    return canonical


async def synchronize_instrument_sources(
    instrument: EvaluationInstrument,
    sources: list[EvaluationSourceFile],
    db: AsyncSession,
) -> None:
    """Atomically update the draft snapshot and revision after a source mutation."""

    general_data = _canonical_source_general_data(instrument.general_data_json, sources)
    instrument.general_data_json = general_data
    instrument.revision += 1

    draft = await db.scalar(
        select(EvaluationDraft).where(EvaluationDraft.instrument_id == instrument.id)
    )
    if draft is not None:
        payload = dict(draft.payload_json)
        payload["general_data"] = general_data
        draft.payload_json = payload
        draft.revision = instrument.revision
    await db.flush()


def serialize_instrument(instrument: EvaluationInstrument) -> InstrumentRead:
    participants = sorted(instrument.participants, key=lambda item: item.sort_order)
    criteria = sorted(instrument.criteria, key=lambda item: item.sort_order)
    participant_by_id = {item.id: item for item in participants}
    criterion_by_id = {item.id: item for item in criteria}
    levels_by_id = {level.id: level for criterion in criteria for level in criterion.levels}

    return InstrumentRead(
        id=instrument.id,
        kind=instrument.kind,
        title=instrument.title,
        roster_id=instrument.roster_id,
        status=instrument.status,
        general_data=instrument.general_data_json,
        settings=instrument.settings_json,
        general_observation=instrument.general_observation,
        revision=instrument.revision,
        archived_at=instrument.archived_at,
        participants=[
            ParticipantRead(
                id=item.id,
                student_id=item.student_id,
                student_name=item.student.full_name,
                internal_code=item.student.internal_code,
                role=item.role,
                team_name=item.team_name,
                sort_order=item.sort_order,
                common_notes=item.common_notes,
                individual_notes=item.individual_notes,
            )
            for item in participants
        ],
        criteria=[
            CriterionRead(
                id=item.id,
                client_key=item.client_key,
                code=item.code,
                title=item.title,
                description=item.description,
                weight=item.weight,
                sort_order=item.sort_order,
                levels=[
                    LevelRead.model_validate(level)
                    for level in sorted(item.levels, key=lambda value: value.sort_order)
                ],
            )
            for item in criteria
        ],
        records=[
            RecordRead(
                id=item.id,
                student_id=participant_by_id[item.participant_id].student_id,
                criterion_id=item.criterion_id,
                criterion_key=criterion_by_id[item.criterion_id].client_key,
                level_id=item.level_id,
                level_key=levels_by_id[item.level_id].client_key if item.level_id else None,
                value=item.value,
                evidence=item.evidence,
                strength=item.strength,
                improvement=item.improvement,
                recommendation=item.recommendation,
                teacher_decision=item.teacher_decision,
                observation=item.observation,
            )
            for item in sorted(
                instrument.records,
                key=lambda value: (
                    participant_by_id[value.participant_id].sort_order,
                    criterion_by_id[value.criterion_id].sort_order,
                ),
            )
        ],
        observations=[
            ObservationRead(
                id=item.id,
                student_id=(
                    participant_by_id[item.participant_id].student_id
                    if item.participant_id
                    else None
                ),
                observed_at=item.observed_at,
                situation=item.situation,
                focus=item.focus,
                objective_facts=item.objective_facts,
                context_factors=item.context_factors,
                interpretation=item.interpretation,
                conclusion=item.conclusion,
                commitments=item.commitments,
                common_to_group=item.common_to_group,
            )
            for item in instrument.observations
        ],
        sources=[
            serialize_source(item, instrument_revision=instrument.revision)
            for item in instrument.sources
        ],
        created_at=instrument.created_at,
        updated_at=instrument.updated_at,
    )


def _excel_text(value: object) -> str:
    if value is None:
        return ""
    # A leading =, +, - or @ can be interpreted as a formula by Excel.  Trim
    # first so control characters (tabs/newlines) cannot be used to hide the
    # dangerous prefix, then force the value to remain literal text.
    text = str(value).replace("*", "").strip()
    if text.startswith(("=", "+", "-", "@")):
        return f"'{text}"
    return text


_CHECKLIST_GENERAL_DATA_ALIASES: dict[str, tuple[str, ...]] = {
    "teacher_name": ("teacher_name", "teacherName"),
    "director_name": ("director_name", "directorName"),
    "institution_name": ("institution_name", "institutionName", "institution"),
    "modality": ("modality",),
    "education_level": ("education_level", "educationLevel", "level"),
    "grade": ("grade",),
    "curricular_area": ("curricular_area", "curricularArea", "area"),
    "activity": ("activity",),
    "date": ("date",),
    "period": ("period",),
}


def _checklist_general_value(data: dict[str, object], canonical_key: str) -> object:
    """Read canonical and legacy checklist keys without breaking saved drafts."""
    for key in _CHECKLIST_GENERAL_DATA_ALIASES.get(canonical_key, (canonical_key,)):
        value = data.get(key)
        if value is not None and str(value).strip():
            return value
    return ""


def build_checklist_xlsx(instrument: EvaluationInstrument) -> bytes:
    if instrument.kind != "checklist":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="La exportación de lista de cotejo solo está disponible para ese instrumento.",
        )
    participants = sorted(instrument.participants, key=lambda item: item.sort_order)
    criteria = sorted(instrument.criteria, key=lambda item: item.sort_order)
    records = {
        (record.participant_id, record.criterion_id): record for record in instrument.records
    }

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Lista de cotejo"
    sheet.sheet_view.showGridLines = False
    title_fill = PatternFill("solid", fgColor="DCEBFF")
    header_fill = PatternFill("solid", fgColor="EDE9FE")
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    last_column = max(4, 4 + len(criteria))
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=last_column)
    title_cell = sheet.cell(1, 1, _excel_text(instrument.title))
    title_cell.font = Font(name="Arial", size=16, bold=True, color="000000")
    title_cell.fill = title_fill
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    sheet.row_dimensions[1].height = 28

    metadata_labels = (
        ("Docente", "teacher_name"),
        ("Director(a)", "director_name"),
        ("Institución", "institution_name"),
        ("Modalidad", "modality"),
        ("Nivel", "education_level"),
        ("Grado / ciclo", "grade"),
        ("Área", "curricular_area"),
        ("Actividad / evidencia", "activity"),
        ("Fecha", "date"),
        ("Periodo", "period"),
    )
    metadata = [
        f"{label}: {_excel_text(_checklist_general_value(instrument.general_data_json, key))}"
        for label, key in metadata_labels
        if _checklist_general_value(instrument.general_data_json, key)
    ]
    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=last_column)
    sheet.cell(2, 1, "  |  ".join(metadata)).font = Font(name="Arial", color="000000")
    sheet.cell(2, 1).alignment = Alignment(wrap_text=True, vertical="top")

    header_row = 4
    headers = ["N.°", "Estudiante", "Código", *[item.code for item in criteria], "Observación"]
    for column, value in enumerate(headers, 1):
        cell = sheet.cell(header_row, column, _excel_text(value))
        cell.font = Font(name="Arial", bold=True, color="000000")
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    value_labels = {"yes": "Sí", "no": "No", "in_progress": "En proceso"}
    for row_index, participant in enumerate(participants, header_row + 1):
        sheet.cell(row_index, 1, row_index - header_row)
        sheet.cell(row_index, 2, _excel_text(participant.student.full_name))
        sheet.cell(row_index, 3, _excel_text(participant.student.internal_code))
        observations: list[str] = []
        for criterion_index, criterion in enumerate(criteria, 4):
            record = records.get((participant.id, criterion.id))
            sheet.cell(
                row_index,
                criterion_index,
                value_labels.get(record.value, "") if record else "",
            )
            if record and record.observation:
                observations.append(
                    f"{_excel_text(criterion.code)}: {_excel_text(record.observation)}"
                )
        if participant.individual_notes:
            observations.append(_excel_text(participant.individual_notes))
        sheet.cell(
            row_index,
            4 + len(criteria),
            _excel_text(" | ".join(observations)),
        )
        for cell in sheet[row_index]:
            cell.font = Font(name="Arial", color="000000")
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    sheet.freeze_panes = "B5"
    sheet.auto_filter.ref = (
        f"A{header_row}:{get_column_letter(last_column)}{header_row + len(participants)}"
    )
    sheet.column_dimensions["A"].width = 7
    sheet.column_dimensions["B"].width = 34
    sheet.column_dimensions["C"].width = 16
    for column in range(4, 4 + len(criteria)):
        sheet.column_dimensions[get_column_letter(column)].width = 15
    sheet.column_dimensions[get_column_letter(4 + len(criteria))].width = 42

    definitions = workbook.create_sheet("Criterios")
    definitions.sheet_view.showGridLines = False
    definition_headers = ["Código", "Criterio", "Descripción completa", "Ponderación"]
    for column, value in enumerate(definition_headers, 1):
        cell = definitions.cell(1, column, value)
        cell.font = Font(name="Arial", bold=True, color="000000")
        cell.fill = header_fill
        cell.border = border
    for row_index, criterion in enumerate(criteria, 2):
        values = (
            criterion.code,
            criterion.title,
            criterion.description,
            criterion.weight,
        )
        for column, value in enumerate(values, 1):
            cell = definitions.cell(row_index, column, _excel_text(value))
            cell.font = Font(name="Arial", color="000000")
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    definitions.freeze_panes = "A2"
    for column, width in enumerate((14, 42, 72, 16), 1):
        definitions.column_dimensions[get_column_letter(column)].width = width

    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def archive_instrument(instrument: EvaluationInstrument) -> None:
    instrument.status = "archived"
    instrument.archived_at = datetime.now(UTC)
    instrument.revision += 1


def restore_instrument(instrument: EvaluationInstrument) -> None:
    instrument.status = "draft"
    instrument.archived_at = None
    instrument.revision += 1
